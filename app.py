from flask import Flask, render_template, request, send_file, Response
from dotenv import load_dotenv
from pydantic import BaseModel
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from langchain.agents import create_tool_calling_agent, AgentExecutor
from langchain_community.tools import WikipediaQueryRun, DuckDuckGoSearchRun
from langchain_community.utilities import WikipediaAPIWrapper
from langchain.tools import Tool
from datetime import datetime
import json
import os
import io
from docx import Document

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Define the output structure
class ResearchResponse(BaseModel):
    topic: str
    summary: str
    sources: list[str]
    tools_used: list[str]

# Create a document in memory
def create_document(data):
    """
    Create a Word document in memory from research data
    """
    # Try to parse as JSON if it's a string
    if isinstance(data, str):
        try:
            data_dict = json.loads(data)
        except json.JSONDecodeError:
            data_dict = {"topic": "Research Topic", "summary": data}
    else:
        # If it's already a dict-like object
        data_dict = data if hasattr(data, "get") else {"topic": "Research Topic", "summary": str(data)}
    
    # Extract information
    topic = data_dict.get("topic", "Unknown Topic")
    summary = data_dict.get("summary", "No summary available")
    sources = data_dict.get("sources", [])
    
    # Create a document in memory
    doc = Document()
    
    # Add title
    doc.add_heading(topic, 0)
    
    # Add summary section
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(summary)
    
    # Add sources section if available
    if sources:
        doc.add_heading('Sources', level=1)
        for source in sources:
            doc.add_paragraph(source, style='List Bullet')
    
    # Add footer with date
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on {datetime.now().strftime('%Y-%m-%d')}"
    
    return doc

# Initialize tools
def setup_agent():
    search = DuckDuckGoSearchRun()
    search_tool = Tool(
        name="search",
        func=search.run,
        description="Search the web for information",
    )

    api_wrapper = WikipediaAPIWrapper(top_k_results=1, doc_content_chars_max=100)
    wiki_tool = Tool(
        name="wiki",
        func=WikipediaQueryRun(api_wrapper=api_wrapper).run,
        description="Search Wikipedia for information",
    )

    # Set up the language model and parser
    llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")
    parser = PydanticOutputParser(pydantic_object=ResearchResponse)

    # Create the prompt template
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
                You are a research assistant that will help generate a research paper.
                Answer the user query and use necessary tools. 
                Wrap the output in this format and provide no other text\n{format_instructions}
                """,
            ),
            ("placeholder", "{chat_history}"),
            ("human", "{query}"),
            ("placeholder", "{agent_scratchpad}"),
        ]
    ).partial(format_instructions=parser.get_format_instructions())

    # Define tools list and create agent
    tools = [search_tool, wiki_tool]
    agent = create_tool_calling_agent(
        llm=llm,
        prompt=prompt,
        tools=tools
    )

    # Create the agent executor
    agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=False)
    
    return agent_executor, parser

# Clean model output
def clean_output(output):
    """Clean the output if it contains markdown code blocks or extra formatting."""
    if isinstance(output, str):
        # Remove markdown code blocks if present
        if output.startswith("```") and output.endswith("```"):
            output = output.strip("`").strip()
            # Extract just the JSON part if it has a language identifier
            if output.startswith("json\n"):
                output = output[5:]
    return output

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/research', methods=['POST'])
def research():
    query = request.form['query']
    
    if not query:
        return render_template('index.html', error="Please enter a research topic")
    
    try:
        # Setup agent and parser
        agent_executor, parser = setup_agent()
        
        # Invoke the agent
        raw_response = agent_executor.invoke({"query": query})
        
        # Process the response
        if "output" in raw_response:
            output_text = clean_output(raw_response["output"])
            
            try:
                # Try to parse the JSON
                structured_response = parser.parse(output_text)
                response_dict = structured_response.model_dump()
                
                return render_template(
                    'results.html',
                    topic=response_dict["topic"],
                    summary=response_dict["summary"],
                    sources=response_dict["sources"],
                    tools_used=response_dict.get("tools_used", []),
                    raw_response=json.dumps(response_dict, indent=2)
                )
                
            except Exception as e:
                # If parsing fails, return raw text
                return render_template(
                    'results.html',
                    topic="Research Results",
                    summary=output_text,
                    sources=[],
                    tools_used=[],
                    error=f"Error parsing results: {str(e)}",
                    raw_response=output_text
                )
        else:
            return render_template('index.html', error="No output found in response")
    
    except Exception as e:
        return render_template('index.html', error=f"Error: {str(e)}")

@app.route('/download')
def download():
    # Get parameters from query string
    topic = request.args.get('topic', 'Research Topic')
    summary = request.args.get('summary', 'No summary available')
    sources = request.args.get('sources', '').split('|')
    
    # Filter empty strings
    sources = [s for s in sources if s]
    
    # Create document
    data = {
        "topic": topic,
        "summary": summary,
        "sources": sources
    }
    
    doc = create_document(data)
    
    # Save to BytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Return file for download
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"research_{topic.replace(' ', '_')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))  # Get the port from the environment variable
    app.run(host='0.0.0.0', port=port, debug=True)